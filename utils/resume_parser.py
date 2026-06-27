# =============================================================================
# utils/resume_parser.py
# =============================================================================
import re, io
import pdfplumber
from docx import Document


def extract_text_from_pdf(file_bytes: bytes) -> str:
    text = ""
    try:
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    text += t + "\n"
    except Exception as e:
        text = f"PDF_ERROR: {e}"
    return text.strip()


def extract_text_from_docx(file_bytes: bytes) -> str:
    text = ""
    try:
        doc = Document(io.BytesIO(file_bytes))
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
    except Exception as e:
        text = f"DOCX_ERROR: {e}"
    return text.strip()


def extract_email(text: str) -> str:
    m = re.search(r'[\w\.-]+@[\w\.-]+\.\w+', text)
    return m.group(0) if m else "Not found"


def extract_phone(text: str) -> str:
    m = re.search(r'(\+?\d[\d\s\-\(\)]{8,15}\d)', text)
    return m.group(0).strip() if m else "Not found"


def extract_linkedin(text: str) -> str:
    m = re.search(r'linkedin\.com/in/[\w\-]+', text, re.IGNORECASE)
    return "https://" + m.group(0) if m else "Not found"


def extract_github(text: str) -> str:
    m = re.search(r'github\.com/[\w\-]+', text, re.IGNORECASE)
    return "https://" + m.group(0) if m else "Not found"


def extract_name(text: str) -> str:
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:4]:
        if 2 <= len(line.split()) <= 5 and not re.search(r'[\d@|/\\:]', line):
            if not any(w in line.lower() for w in ['resume','cv','curriculum','objective','summary']):
                return line
    return lines[0] if lines else "Not detected"


def parse_resume(uploaded_file) -> dict:
    file_bytes = uploaded_file.read()
    name = uploaded_file.name.lower()
    if name.endswith(".pdf"):
        text = extract_text_from_pdf(file_bytes)
    elif name.endswith(".docx"):
        text = extract_text_from_docx(file_bytes)
    else:
        text = ""
    return {
        "raw_text":   text,
        "name":       extract_name(text),
        "email":      extract_email(text),
        "phone":      extract_phone(text),
        "linkedin":   extract_linkedin(text),
        "github":     extract_github(text),
        "word_count": len(text.split()),
        "char_count": len(text),
        "filename":   uploaded_file.name,
    }
